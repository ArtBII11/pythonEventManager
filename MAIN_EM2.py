import sys
import csv
import os
import smtplib
from email.message import EmailMessage
import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
import os
import docx


CSV_FILE_MEET = "Meetings.csv"
CSV_FILE = "users.csv"
text_guest = ""
text_member = ""


def attach_file_doc(id_button):
    global text_guest
    global text_member
    file_formats = [
        ("Документы (*.txt, *.docx)", "*.txt *.docx"),
        ("Текстовые файлы (*.txt)", "*.txt"),
        ("Документы Word (*.docx)", "*.docx"),
    ]

    file_path = filedialog.askopenfilename(
        title="Выберите файл для прикрепления",
        filetypes=file_formats,
    )

    if file_path:
        print(f"Файл успешно прикреплен: {file_path}")

        # Получаем расширение файла (например, '.txt' или '.docx')
        _, extension = os.path.splitext(file_path.lower())

        try:
            # ЕСЛИ ЭТО ФАЙЛ WORD (.docx)
            if extension == ".docx":
                if id_button == "member":
                    doc = docx.Document(file_path)
                    # Собираем все абзацы из файла Word в одну текстовую переменную,
                    # разделяя их обычными переносами строк \n
                    text_member = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif id_button == "guest":
                    doc = docx.Document(file_path)
                    # Собираем все абзацы из файла Word в одну текстовую переменную,
                    # разделяя их обычными переносами строк \n
                    text_guest = "\n".join([paragraph.text for paragraph in doc.paragraphs])                    

            # ЕСЛИ ЭТО ОБЫЧНЫЙ ТЕКСТ (.txt)
            else:
                if id_button =="member":
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_member = f.read()
                else:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_guest = f.read()

            # Проверяем, удалось ли вытащить текст
            if text_member.strip() and id_button == "member":
                messagebox.showinfo(
                    "Успех", f"Файл прочитан! Найдено {len(text_member)} символов."
                )
            elif text_guest.strip() and id_button == "guest":
                messagebox.showinfo(
                    "Успех", f"Файл прочитан! Найдено {len(text_guest)} символов."
                )
            else:
                messagebox.showwarning(
                    "Внимание", "Файл открыт, но текста внутри не обнаружено."
                )                

        except Exception as e:
            messagebox.showerror(
                "Ошибка",
                f"Не удалось прочитать файл. Возможно, он открыт в другой программе.\nОшибка: {e}",
            )


def load_data_from_csv(csv_file):
    """Функция загрузки данных из CSV-файла."""
    if not os.path.exists(csv_file):
        create_test_csv(csv_file)

    data = []
    try:
        with open(csv_file, mode="r", encoding="utf-8-sig") as f:
            reader = csv.reader(f, delimiter=";")
            next(reader)  # Пропускаем заголовок

            for row in reader:
                if row:  # Проверка на пустые строки
                    data.append(row)
    except Exception as e:
        messagebox.showerror(
            "Ошибка", f"Не удалось прочитать файл {CSV_FILE}:\n{e}"
        )
    return data


def create_test_csv(csv_file):
    """Создает тестовый файл, если его не существует."""
    if csv_file == "users.csv":
        test_data = [
            ["id", "name", "email", "role"],
            ["1", "Алексей", "artyombirulya@gmail.com", "Админ"],  # Ваш email для теста
            ["2", "Мария", "maria@example.com", "Пользователь"],
            ["3", "Иван", "ivan@example.com", "Модератор"],
        ]
    elif csv_file == "Meetings.csv":
        test_data = [
                ["name_of_meet", "end_date", "start_date"],
                ["Met", "02.04.26", "02.04.26"],  # Ваш email для теста
                ["et2", "02.04.26", "02.04.26"],
                ["IRL", "02.04.26", "02.04.26"],
            ]
    elif csv_file == "guests.csv":
        test_data = [
                    ["id", "name", "email"],
                    ["1", "Артём", "artyombirulya@gmail.com"],  # Ваш email для теста
                    ["2", "Мария", "maria@example.com"],
                    ["3", "Иван", "ivan@example.com"],
                ]
    with open(csv_file, mode="w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerows(test_data)



def send_emails():
    global text_member
    global text_guest
    """Функция рассылки писем по всему списку из таблицы."""
    all_items_members = tree_of_users.get_children()
    all_items_guests = tree_of_guest.get_children()

    if not all_items_members:
        messagebox.showwarning("Внимание", "Список участникоы пуст!")
        return
    elif not all_items_guests:
        messagebox.showwarning("Внимание", "Список гостей пуст!")
        return
    elif not all_items_guests and not all_items_members:
        messagebox.showwarning("Внимание", "Все списки гостей и участников пустые. Ты тупой?")
        return

    sent_count = 0

    all_items = []

    for item in all_items_members:
        all_items.append((item,"member",tree_of_users))
    for item in all_items_guests:
            all_items.append((item,"guest",tree_of_guest))



    # Проходим циклом по каждой строке таблицы
    for item,role,tree in all_items:
        values = tree.item(item, "values")
        if len(values)<3:
            continue


        user_name = values[1]
        user_email = values[2]

        if role == "member":
            current_text = text_member
            sub_prex = "участник"
        elif role == "guest":
            current_text = text_guest
            sub_prex = "гость"

        print(f"Отправка на {user_email}: Hello, {user_name}!")

        # НАСТРОЙКИ ЯНДЕКСА
        SENDER = "artmbirulya@yandex.com"  # Изменили .com на .ru
        PASSWORD = "fiwvbwzisbfnincq"  # Ваш 16-значный пароль приложения
        RECIPIENT = user_email

        custom_text = current_text.format(
            username = user_name,
            meet_name = entry_name.get(),
            meet_time = entry_time.get(),
            meet_date = entry_data.get()
        )
        
        
        # Создание письма
        msg = EmailMessage()
        msg["Subject"] = f" Приглашение на мероприятие {entry_name.get()}"
        print(entry_name.get())
        msg["From"] = SENDER
        msg["To"] = RECIPIENT
        msg.set_content(
            f"Привет, {user_name}! Вы наш {sub_prex} данного мероприятия, дата мероприятия: {entry_data.get()} и время: {entry_time.get()}.\n Дальнейшие инструкции здесь: \n {custom_text}"
        )
        # Отправка через сервер Яндекса
        try:
            with smtplib.SMTP_SSL("smtp.yandex.ru", 465) as server:
                server.set_debuglevel(1)  # Логи включаются ДО авторизации
                server.login(SENDER, PASSWORD)
                server.send_message(msg)

            print(f"Письмо для {user_name} успешно отправлено!")
            sent_count += 1  # Считаем только успешные

        except Exception as e:
            print(f"Произошла ошибка при отправке на {user_email}: {e}")
            messagebox.showerror(
                "Произошла ошибка при отправке",
                f"Не удалось отправить письмо для {user_name} ({user_email}):\n{e}",
            )

    # ОКНО УВЕДОМЛЕНИЯ (Вынесено из цикла наружу)
    if sent_count > 0:
        messagebox.showinfo(
            "Рассылка завершена",
            f"Успешно отправлено писем: {sent_count}.",
        )


placeholder_map = {}

def enter_function_input(event):
    text_placeholder = placeholder_map.get(event.widget)
    if event.widget.get() != text_placeholder and event.widget.get() != "":
        event.widget.config(state="readonly")

#Гуишки для окна
root = tk.Tk()
root.title("Event manager(Editing)")
#root.resizable(False, False)
root.geometry("750x900")

container_meetings = ttk.Frame(root)
container_meetings.grid(column=1,row=5,padx=10,sticky="w")

scrollbar_meetings = ttk.Scrollbar(container_meetings, orient=tk.VERTICAL)#----Скролбар
scrollbar_meetings.grid(column=2,sticky="w",row=5,ipady=100)

container_users = ttk.Frame(root)
container_users.grid(column=1,row=3,padx=10,sticky="w")

scrollbar_users = ttk.Scrollbar(container_users, orient=tk.VERTICAL)#----Скролбар
scrollbar_users.grid(column=2,sticky="w",row=3,ipady=100)

container_guest = ttk.Frame(root)
container_guest.grid(column=1,row=4,padx=10,sticky="w")

scrollbar_guest = ttk.Scrollbar(container_guest, orient=tk.VERTICAL)#----Скролбар
scrollbar_guest.grid(column=2,sticky="w",row=4,ipady=100)

# Создаем таблицу (Treeview)

columns_meetings = ("name_of_meet", "date", "time")
tree_of_meetings = ttk.Treeview(container_meetings, columns=columns_meetings, show="headings", yscrollcommand=scrollbar_meetings.set)
scrollbar_meetings.config(command=tree_of_meetings.yview)

columns_users = ("id", "name", "email", "role")
tree_of_users = ttk.Treeview(container_users, columns=columns_users, show="headings",yscrollcommand=scrollbar_users.set)
scrollbar_users.config(command=tree_of_users.yview)

tree_of_meetings.heading("name_of_meet", text="Название мероприятия")
tree_of_meetings.heading("date", text="Дата мероприятия")
tree_of_meetings.heading("time", text="Время мероприятия")

# Задаем размеры колонок
tree_of_meetings.column("name_of_meet", width=120, anchor=tk.CENTER)
tree_of_meetings.column("date", width=120)
tree_of_meetings.column("time", width=120)

meet_data_t = load_data_from_csv("Meetings.csv")
for meet in meet_data_t:
    tree_of_meetings.insert("", tk.END, values=meet)

tree_of_meetings.grid(column=1,row=5,ipadx=0,ipady=0)

# Задаем заголовки колонок
tree_of_users.heading("id", text="ID")
tree_of_users.heading("name", text="Имя")
tree_of_users.heading("email", text="Email")
tree_of_users.heading("role", text="Роль")

# Задаем размеры колонок
tree_of_users.column("id", width=40, anchor=tk.CENTER)
tree_of_users.column("name", width=120)
tree_of_users.column("email", width=180)
tree_of_users.column("role", width=100)

# Загружаем данные из файла и вставляем в таблицу
users_data = load_data_from_csv("users.csv")
for user in users_data:
    tree_of_users.insert("", tk.END, values=user)

tree_of_users.grid(column=1,row=3,ipadx=0,ipady=0,pady=0)

columns_guest = ("id", "name", "email")
tree_of_guest = ttk.Treeview(container_guest, columns=columns_guest, show="headings",yscrollcommand=scrollbar_guest.set)
scrollbar_guest.config(command=tree_of_users.yview)

tree_of_guest.heading("id", text="ID")
tree_of_guest.heading("name", text="Имя")
tree_of_guest.heading("email", text="Email")


# Задаем размеры колонок
tree_of_guest.column("id", width=40, anchor=tk.CENTER)
tree_of_guest.column("name", width=120)
tree_of_guest.column("email", width=180)

# Загружаем данные из файла и вставляем в таблицу
guest_data = load_data_from_csv("guests.csv")
for guest in guest_data:
    tree_of_guest.insert("", tk.END, values=guest)

tree_of_guest.grid(column=1,row=4,ipadx=0,ipady=0,pady=0)

style = ttk.Style()
style.configure("My.TButton",font=("Arial", 12, "bold"),)

btn_send = ttk.Button(
    root, text="Разослать всем пользователям", command=send_emails,style="My.TButton")
btn_send.grid(column=1, row=6, columnspan=3, pady=15, ipady=10)

btn_attach_member = ttk.Button(root,text="Прикрепить шаблон участника",command=lambda:attach_file_doc("member"),style="My.TButton")
btn_attach_member .grid(column=2, row=3, pady=15, ipady=10)

btn_attach_guest = ttk.Button(root,text="Прикрепить шаблон гостя",command=lambda:attach_file_doc("guest"),style="My.TButton")
btn_attach_guest.grid(column=2, row=4, pady=15, ipady=10)

def off_entry_mouse(event):
    current_entry = event.widget
    text_placeholder = placeholder_map.get(current_entry)
    if current_entry.get() == "":
        current_entry.insert(0, text_placeholder)

def on_entry_click(event):
    current_entry = event.widget 
    text_placeholder = placeholder_map.get(current_entry)
    if current_entry.get() == text_placeholder:
        current_entry.delete(0, tk.END)


entry_name = ttk.Entry(root,font=("Arial", 13))
entry_name.grid(column=2,row=0,ipady=10,ipadx=20,pady=20,padx=(0,15))
placeholder_map[entry_name] = "Введите сюда название..."
entry_name.insert(0, "Введите сюда название...")


entry_name.bind("<FocusIn>", on_entry_click)
entry_name.bind("<FocusOut>",off_entry_mouse)
entry_name.bind("<Return>",enter_function_input)




entry_time = ttk.Entry(root)
entry_time.grid(column=1,row=0,pady=(80,5),ipady=5,sticky="w",padx=10)
entry_time.insert(0,"Время...")
placeholder_map[entry_time] = "Время..."
entry_time.bind("<FocusIn>", on_entry_click)
entry_time.bind("<FocusOut>",off_entry_mouse)
entry_time.bind("<Return>",enter_function_input)


entry_data = ttk.Entry(root)
entry_data.grid(column=1,row=1,pady=(5,0),ipady=5,sticky="w",padx=10)
entry_data.insert(0,"Дата...")
placeholder_map[entry_data] = "Дата..."
entry_data.bind("<FocusIn>", on_entry_click)
entry_data.bind("<FocusOut>",off_entry_mouse)
entry_data.bind("<Return>",enter_function_input)



# label_title = tk.Label(root, text="Список запланированных мероприятий",font=("Arial",12,"bold"))
# label_title.pack(pady=(20,0))

# container = ttk.Frame(root,height=200,width=600)#----Контейнер для схемы
# container.pack_propagate(False)
# container.pack(padx=20,pady=10)




# columns = ("name", "email", "role")
# tree = ttk.Treeview(
#     container, columns=columns, show="headings", yscrollcommand=scrollbar.set
# )
# scrollbar.config(command=tree.yview)

# # Задаем заголовки колонок

# tree.heading("name", text="Название мероприятия")
# tree.heading("email", text="Конец мероприятия")
# tree.heading("role", text="Начало мероприятия")

# # Задаем размеры колонок
# tree.column("name", width=120, anchor=tk.CENTER)
# tree.column("email", width=180)
# tree.column("role", width=100)

# users_data = load_data_from_csv()
# for user in users_data:
#     tree.insert("", tk.END, values=user)

# tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

root.mainloop()